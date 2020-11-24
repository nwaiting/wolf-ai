import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
import datetime


def get_equation(ss, rd, mn, min_cal_num=5):
    if ss == '+':
        if rd == 2:
            return "{:2} + {:2} = (       )".format(random.randint(min_cal_num, mn), random.randint(min_cal_num, mn))
        else:
            first_num = random.randint(min_cal_num, mn)
            second_num = random.randint(0, first_num)
            if rd == 0:
                return "(       ) + {:2} = {:2}".format(second_num, first_num)
            else:
                return "{:2} + (       ) = {:2}".format(second_num, first_num)
    if ss == '-':
        if rd == 0:
            return "(       ) - {:2} = {:2}".format(random.randint(min_cal_num, mn), random.randint(min_cal_num, mn))
        else:
            first_num = random.randint(min_cal_num, mn)
            second_num = random.randint(0, first_num)
            if rd == 0:
                return "{:2} - {:2} = (       )".format(first_num, second_num)
            else:
                return "{:2} - (       ) = {:2}".format(first_num, second_num)


def main(max_line, max_num, min_cal_num, file_name):
    d = docx.Document()
    lines_per_page = 25
    for _ in range(int(max_line/lines_per_page)):
        # p = d.add_paragraph('{:35}\t{:35}\t{:35}'.format('日期:', '用时:', '打分:'))
        p = d.add_paragraph('')
        p.add_run('{:55}'.format('日期:')).bold = True
        p.add_run('{:55}'.format('用时:')).bold = True
        p.add_run('{:55}'.format('打分:')).bold = True
        # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        symbol_select_list = ['+', '-']
        for i in range(lines_per_page):
            equations = []
            for _ in range(4):
                symbol_select = random.randint(0, 1)
                result_direction = random.randint(0, 2)
                res = get_equation(symbol_select_list[symbol_select], result_direction, max_num, min_cal_num)
                equations.append(res)
            res = "{:25}\t{:25}\t{:25}\t{:25}".format(equations[0], equations[1], equations[2], equations[3])
            d.add_paragraph(res)
        if _ < int(max_line/lines_per_page) - 2:
            d.add_page_break()
    d.save(file_name)


# 100行每一行3个题目，20以内的，最小的数字是8以上的
m_num = 20
main(200, m_num, 8, r'C:\Users\xujie\Desktop\{}以内的加减逆运算.docx'.format(m_num))



